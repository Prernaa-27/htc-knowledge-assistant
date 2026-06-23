import os
import sys
from pathlib import Path

# Ensure the project root is on sys.path so backend imports resolve correctly
ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

from docx import Document
from backend import parser
from backend.embeddings import generate_embeddings
from backend.vectorstore import create_index, save_index, save_chunks
from backend.rag_pipeline import run_rag

sample_path = Path('data/processed_docs') / 'sample_test.docx'
sample_path.parent.mkdir(parents=True, exist_ok=True)
doc = Document()
doc.add_paragraph('Acme Corporation policy states that employees must reset their passwords every 90 days.')
doc.add_paragraph('The refund policy allows returns within 30 days with receipt.')
doc.save(sample_path)

print('Sample DOCX written to', sample_path)

chunks = parser.parse_document(str(sample_path), chunk_size=200, overlap=40)
print('Chunks:')
for i, chunk in enumerate(chunks, 1):
    print(f'  {i}: {chunk}')

vectors = generate_embeddings(chunks)
print('Embeddings shape:', vectors.shape)
index = create_index(vectors)
print('Index total vectors:', index.ntotal)

save_index_path = save_index()
save_chunks_path = save_chunks(chunks)
print('Saved index to', save_index_path)
print('Saved chunks to', save_chunks_path)

question = 'What is Acme Corporation password reset policy?'
result = run_rag(question, chunks, top_k=2)
print('\nRAG result:')
print('Answer:', result.get('answer'))
print('Citations:', result.get('citations'))
print('Chunk IDs:', result.get('chunk_ids'))
